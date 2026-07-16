"""
Génération du document Word (.docx) du Bon de Livraison (BL) — réplique la mise
en page du PDF (voir generate_bl.py) : gabarit générique ou gabarit SARL INDUREX
(en-tête vert, bloc référence, footer RC/NIF + badges ISO) selon le récupérateur.
"""
import io
from docx import Document
from docx.shared import Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

from apps.bc.generate_bc_word import (
    _shade_cell, _cell_borders, _set_col_widths, _set_run, _cell_lines, _kv_rows,
    _add_picture_safe, _zero_spacing, _doc_p, _GENERIC_GREEN, _INDUREX_GREEN,
    _INDUREX_GREEN_HEX, _GENERIC_GREEN_HEX, _WHITE, COL,
)
from .generate_bl import _recuperateur_info, _destinataire_info, _fmt_date, _fmt_montant, _fmt_qte, _is_indurex

_INDUREX_NOM    = 'SARL INDUREX'
_INDUREX_SLOGAN = 'INDUSTRIAL WAST RECOVERY AND VALORIZATION'
_MODE_LIV_ABBR  = {'ENLEVEMENT': 'ENLEV', 'LIVRAISON': 'LIVR'}


def generate_bl_docx(data: dict) -> bytes:
    rec  = _recuperateur_info(data)
    dest = _destinataire_info(data)
    if _is_indurex(rec):
        return _generate_bl_docx_indurex(data, rec, dest)
    return _generate_bl_docx_generique(data, rec, dest)


# ── Gabarit générique ───────────────────────────────────────────────────────────

def _generate_bl_docx_generique(data: dict, rec: dict, dest: dict) -> bytes:
    def v(key, default=''):
        val = data.get(key, default)
        return str(val) if val not in (None, '') else default

    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin   = Cm(1.2)
        section.bottom_margin = Cm(1.2)

    entete = doc.add_table(rows=1, cols=2)
    _set_col_widths(entete, [2.5, COL - 2.5])
    if rec['logo_path']:
        _add_picture_safe(entete.rows[0].cells[0], rec['logo_path'], 2.2, 2.2)
    _cell_lines(entete.rows[0].cells[1], [{
        'text': (rec['nom'] or '').upper(), 'size': 20, 'bold': True, 'italic': True, 'color': _GENERIC_GREEN,
    }])

    if rec['agrement_num']:
        _doc_p(doc, f"Agrément N° {rec['agrement_num']} du {rec['agrement_date']}", size=9)
    adresse_ligne = ' '.join(filter(None, [rec['adresse'], rec['code_postal']]))
    if adresse_ligne:
        _doc_p(doc, adresse_ligne, size=9)

    id_table = doc.add_table(rows=2, cols=2)
    _set_col_widths(id_table, [COL / 2, COL / 2])
    _cell_lines(id_table.rows[0].cells[0], [{'text': f"RC {rec['rc']}", 'size': 9}])
    _cell_lines(id_table.rows[0].cells[1], [{'text': f"NIF {rec['nif']}", 'size': 9}])
    _cell_lines(id_table.rows[1].cells[0], [{'text': f"NA {rec['na']}", 'size': 9}])
    _cell_lines(id_table.rows[1].cells[1], [{'text': f"NIS {rec['nis']}", 'size': 9}])
    _doc_p(doc)

    _doc_p(doc, f"{rec['commune']} le : {_fmt_date(v('date_livraison'))}",
           align=WD_ALIGN_PARAGRAPH.RIGHT, size=9.5)
    _doc_p(doc)

    p1 = _doc_p(doc, 'Nom de Client : ', size=9.5)
    _set_run(p1.add_run(dest['nom']), size=9.5, bold=True)
    p2 = _doc_p(doc, 'Adresse : ', size=9.5)
    _set_run(p2.add_run(dest['adresse']), size=9.5, bold=True)
    _doc_p(doc)

    titre_tbl = doc.add_table(rows=1, cols=1)
    _set_col_widths(titre_tbl, [8])
    _cell_borders(titre_tbl.rows[0].cells[0])
    _cell_lines(titre_tbl.rows[0].cells[0], [{
        'text': 'Bon de livraison', 'size': 13, 'bold': True, 'italic': True, 'align': WD_ALIGN_PARAGRAPH.CENTER,
    }])
    _doc_p(doc)

    lignes  = data.get('lignes') or []
    headers = ['N°', 'Description (Nature des déchets)', 'Quantités', 'Unités', 'Stockage']
    col_w   = [1.3, 7.7, 2.5, 2.5, 3]
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, col_w)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, _GENERIC_GREEN_HEX)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cell.paragraphs[0].add_run(h), size=9, bold=True, color=_WHITE)
    for i, l in enumerate(lignes, start=1):
        row = tbl.add_row().cells
        vals = [str(i), str(l.get('description', '')), str(l.get('quantite', '')),
                 str(l.get('unite', 'KG')), str(l.get('stockage', ''))]
        for j, val in enumerate(vals):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run(row[j].paragraphs[0].add_run(val), size=9)

    _doc_p(doc)
    _doc_p(doc, f"Nom de chauffeur : {v('chauffeur_nom')}", size=9.5)
    _doc_p(doc, f"Immatriculation de camion : {v('camion_immatriculation')}", size=9.5)
    _doc_p(doc)

    if rec['cachet_path'] or rec['signature_path']:
        sign_p = _doc_p(doc, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if rec['cachet_path']:
            _add_picture_safe(sign_p, rec['cachet_path'], 2.8)
        if rec['signature_path']:
            sign_p.add_run('   ')
            _add_picture_safe(sign_p, rec['signature_path'], 3, 1.6)
    _doc_p(doc, 'Le Gérant', align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if rec['responsable']:
        _doc_p(doc, rec['responsable'], align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── Gabarit SARL INDUREX ────────────────────────────────────────────────────────

def _generate_bl_docx_indurex(data: dict, rec: dict, dest: dict) -> bytes:
    def v(key, default=''):
        val = data.get(key, default)
        return str(val) if val not in (None, '') else default

    lignes = data.get('lignes') or []

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1)
        section.bottom_margin = Cm(1.5)

    entete = doc.add_table(rows=1, cols=3)
    _set_col_widths(entete, [2.3, COL - 2.3 - 6.7, 6.7])
    if rec['logo_path']:
        _add_picture_safe(entete.rows[0].cells[0], rec['logo_path'], 2, 2)

    _cell_lines(entete.rows[0].cells[1], [
        {'text': _INDUREX_NOM, 'size': 20, 'bold': True, 'color': _INDUREX_GREEN},
        {'text': _INDUREX_SLOGAN, 'size': 9.5, 'bold': True, 'color': _INDUREX_GREEN},
    ])

    ref_cell = entete.rows[0].cells[2]
    _cell_borders(ref_cell)
    ref_tbl = ref_cell.add_table(rows=4, cols=2)
    _set_col_widths(ref_tbl, [2.6, 4.1])
    _kv_rows(ref_tbl, [
        ('Référence', v('numero')),
        ('Date',      _fmt_date(v('date_livraison'))),
        ('Montant',   _fmt_montant(data.get('montant_reference') or 0)),
        ('Mode Liv',  _MODE_LIV_ABBR.get(v('mode_livraison'), v('mode_livraison'))),
    ], label_size=9.5, value_size=9.5)

    _doc_p(doc)

    titre_tbl = doc.add_table(rows=1, cols=1)
    _set_col_widths(titre_tbl, [COL])
    titre_cell = titre_tbl.rows[0].cells[0]
    _shade_cell(titre_cell, _INDUREX_GREEN_HEX)
    _cell_lines(titre_cell, [{
        'text': f"Bon Livraison N°: {v('numero')}", 'size': 16, 'bold': True, 'color': _WHITE,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    }])

    _doc_p(doc)

    client_lignes = [
        ('Réf Client',   v('ref_client')),
        ('N° RC',        v('client_rc')),
        ('NIF',          v('client_nif')),
        ('N° Article',   v('client_numero_article')),
        ('N° I.S',       v('client_nis')),
        ('Tél',          v('client_telephone')),
        ('Fax',          v('client_fax')),
        ('Email',        v('client_email')),
        ('Pièces Liées', v('pieces_liees')),
    ]
    bloc_client = doc.add_table(rows=1, cols=2)
    _set_col_widths(bloc_client, [8.5, 8.5])
    gauche_cell = bloc_client.rows[0].cells[0]
    gauche_tbl  = gauche_cell.add_table(rows=len(client_lignes), cols=2)
    _set_col_widths(gauche_tbl, [2.6, 5.9])
    _kv_rows(gauche_tbl, client_lignes, label_size=10.5, value_size=10.5)

    # Nichée dans sa propre table (cf. generate_bc_word.py) pour que le cadre
    # n'épouse que son propre contenu, sans s'étirer sur les 9 lignes de gauche.
    droite_outer = bloc_client.rows[0].cells[1]
    droite_tbl = droite_outer.add_table(rows=1, cols=1)
    _set_col_widths(droite_tbl, [8.5])
    droite_cell = droite_tbl.rows[0].cells[0]
    _cell_borders(droite_cell)
    droite_lines = [{'text': dest['nom'], 'size': 10.5, 'bold': True}]
    for ligne_adresse in (dest['adresse'] or '').split('\n'):
        if ligne_adresse.strip():
            droite_lines.append({'text': ligne_adresse.strip(), 'size': 10.5})
    _cell_lines(droite_cell, droite_lines)

    _doc_p(doc)

    col_w   = [2.8, 9.2, 2.3, 2.7]
    headers = ['Réf Article', 'Désignation', 'Unité', 'Quantité']
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, col_w)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, _INDUREX_GREEN_HEX)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=_WHITE)
    for l in lignes:
        row = tbl.add_row().cells
        vals = [str(l.get('ref_article', '')), str(l.get('description', '')),
                str(l.get('unite', 'KG')), _fmt_qte(l.get('quantite'))]
        for j, val in enumerate(vals):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _set_run(row[j].paragraphs[0].add_run(val), size=10)

    if rec['cachet_path'] or rec['signature_path']:
        _doc_p(doc)
        sign_p = _doc_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        if rec['cachet_path']:
            _add_picture_safe(sign_p, rec['cachet_path'], 4)
        if rec['signature_path']:
            sign_p.add_run('   ')
            _add_picture_safe(sign_p, rec['signature_path'], 4.5, 2.4)

    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ''
    _zero_spacing(footer_p)
    _set_run(footer_p.add_run(
        f"RC: {rec['rc']}  NIF: {rec['nif']}  Al: {rec['na']}  NIS: {rec['nis']}"
    ), size=7)
    if rec['adresse']:
        p_adr = _zero_spacing(footer.add_paragraph())
        _set_run(p_adr.add_run(rec['adresse']), size=7)
    footer_extra = '  '.join(filter(None, [
        rec['commune'],
        f"Email: {rec['email']}" if rec['email'] else '',
        f"Tél: {rec['telephone']}" if rec['telephone'] else '',
        f"Fax: {rec['fax']}" if rec['fax'] else '',
    ]))
    if footer_extra:
        p_extra = _zero_spacing(footer.add_paragraph())
        _set_run(p_extra.add_run(footer_extra), size=7)

    iso_paths = [p for p in (rec.get('iso_9001_path'), rec.get('iso_14001_path'), rec.get('iso_45001_path')) if p]
    if iso_paths:
        p_iso = _zero_spacing(footer.add_paragraph())
        p_iso.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for path in iso_paths:
            _add_picture_safe(p_iso, path, 1.5)
            p_iso.add_run('  ')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
