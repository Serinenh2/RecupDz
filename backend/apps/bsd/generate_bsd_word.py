"""
Génération du document Word (.docx) du BSD — Bordereau de Suivi des Déchets.
Contenu équivalent au PDF, présenté de façon simple et imprimable sous Word.
"""
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _titre(doc, texte):
    p = doc.add_heading(texte, level=2)
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER


def _section(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(12)


def _champ(doc, label, valeur):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} : ")
    r1.bold = True
    p.add_run(str(valeur) if valeur else '—')


def _fmt_date(iso):
    if not iso:
        return ''
    parts = str(iso).split('-')
    if len(parts) != 3:
        return str(iso)
    y, m, d = parts
    return f"{d}/{m}/{y}"


def generate_bsd_docx(data: dict) -> bytes:
    doc = Document()
    doc.add_heading('BORDEREAU DE SUIVI DES DÉCHETS (BSD)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    sous = doc.add_paragraph('Conforme au Décret exécutif n°06-104 du 28 février 2006')
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _champ(doc, 'N° BSD', data.get('numero'))
    _champ(doc, 'Date d\'émission', _fmt_date(data.get('date_emission')))
    _champ(doc, 'Statut', data.get('statut_display') or data.get('statut'))

    doc.add_paragraph()
    _section(doc, '1 — Générateur des déchets')
    _champ(doc, 'Raison sociale', data.get('generateur_nom'))
    _champ(doc, 'Adresse', data.get('generateur_adresse'))

    doc.add_paragraph()
    _section(doc, '2 — Identification du déchet')
    _champ(doc, 'Code déchet', data.get('code_dechet'))
    _champ(doc, 'Désignation', data.get('designation'))
    _champ(doc, 'Classe', data.get('classe'))

    doc.add_paragraph()
    _section(doc, '3 — Quantité et conditionnement')
    _champ(doc, 'Quantité', f"{data.get('quantite','')} {data.get('unite_display') or data.get('unite','')}")
    _champ(doc, 'Emballage / Conditionnement', data.get('emballage'))

    doc.add_paragraph()
    _section(doc, '4 — Transporteur')
    _champ(doc, 'Société', data.get('transporteur_nom'))
    _champ(doc, 'Véhicule', data.get('transporteur_vehicule'))

    doc.add_paragraph()
    _section(doc, '5 — Destination finale')
    _champ(doc, 'Destinataire / Récepteur', data.get('recepteur_nom') or data.get('destination_nom'))
    _champ(doc, 'Type de traitement', data.get('type_traitement'))
    _champ(doc, 'Date de réception', _fmt_date(data.get('date_reception')))

    if data.get('observations'):
        doc.add_paragraph()
        _section(doc, 'Observations')
        doc.add_paragraph(data.get('observations'))

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
