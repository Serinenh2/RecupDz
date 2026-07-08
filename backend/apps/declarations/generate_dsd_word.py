"""
Génération du document Word (.docx) de la DSD — Déclaration des Déchets
Spéciaux Dangereux. Contenu équivalent au PDF, présenté de façon simple et
imprimable sous Word — conforme au Décret exécutif n°05-315 du 10/09/2005.
"""
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH


def _section(doc, texte):
    p = doc.add_paragraph()
    run = p.add_run(texte)
    run.bold = True
    run.font.size = Pt(12)


def _champ(doc, label, valeur):
    p = doc.add_paragraph()
    r1 = p.add_run(f"{label} : ")
    r1.bold = True
    p.add_run(str(valeur) if valeur else '—')


def _fmt_date(iso):
    if not iso:
        return ''
    parts = str(iso).split('-')
    if len(parts) != 3:
        return str(iso)
    y, m, d = parts
    return f"{d}/{m}/{y}"


def generate_dsd_docx(data: dict) -> bytes:
    doc = Document()
    doc.add_heading('DÉCLARATION DES DÉCHETS SPÉCIAUX DANGEREUX (DSD)', level=1).alignment = WD_ALIGN_PARAGRAPH.CENTER
    sous = doc.add_paragraph('Décret exécutif n°05-315 du 10 septembre 2005')
    sous.alignment = WD_ALIGN_PARAGRAPH.CENTER

    doc.add_paragraph()
    _section(doc, 'Identification du générateur')
    _champ(doc, 'Année', data.get('annee'))
    _champ(doc, 'Date de transmission', _fmt_date(data.get('date_transmission')))
    _champ(doc, 'Statut juridique', data.get('statut_juridique'))
    _champ(doc, 'Dénomination', data.get('denomination'))
    _champ(doc, 'Siège social', data.get('siege_social'))
    _champ(doc, 'Domaine d\'activité', data.get('domaine_activite'))
    _champ(doc, 'Certification', data.get('certification'))
    _champ(doc, 'Responsable déchets', data.get('responsable_dechets'))

    doc.add_paragraph()
    _section(doc, 'A — Nature, quantité et caractéristiques des déchets spéciaux dangereux générés')
    _champ(doc, 'Matière première', data.get('matiere_premiere'))
    _champ(doc, 'Consistance', data.get('consistance'))
    _champ(doc, 'Code du déchet', data.get('code_dechet'))
    _champ(doc, 'Dénomination du déchet', data.get('denomination_dechet'))
    _champ(doc, 'Autres précisions', data.get('autres_precisions'))
    _champ(doc, 'Quantité générée (t/an)', data.get('quantite_generee'))
    _champ(doc, 'Composition chimique', data.get('composition_chimique'))
    _champ(doc, 'Critère de dangerosité', data.get('critere_dangerosite'))
    _champ(doc, 'Stockage temporaire (t/an)', data.get('stockage_temporaire_qte'))
    _champ(doc, 'Stockage permanent (t/an)', data.get('stockage_permanent_qte'))
    _champ(doc, 'Modalités de stockage', data.get('modalites_stockage'))

    doc.add_paragraph()
    _section(doc, 'B — Modes de traitement')
    _champ(doc, 'Modalités de gestion', data.get('modalites_gestion'))
    _champ(doc, 'Modalités de contrôle', data.get('modalites_controle'))
    _champ(doc, 'Modalités d\'élimination', data.get('modalites_elimination'))
    _champ(doc, 'Types d\'installation', data.get('types_installation'))
    _champ(doc, 'Types de traitement', data.get('types_traitement'))
    _champ(doc, 'Quantités traitées (t/an)', data.get('quantites_traitees'))
    _champ(doc, 'Rendement', data.get('rendement_traitement'))

    doc.add_paragraph()
    _section(doc, 'C — Mesures prises et à prévoir pour éviter la production des déchets spéciaux dangereux')
    _champ(doc, 'Réutilisation (t/an)', data.get('reutilisation_qte'))
    _champ(doc, 'Recyclage (t/an)', data.get('recyclage_qte'))
    _champ(doc, 'Valorisation (t/an)', data.get('valorisation_qte'))
    _champ(doc, 'Élimination (t/an)', data.get('elimination_qte'))

    mesures = [
        ('1 — Techniques de minimisation', 'mesures_min_prises', 'mesures_min_envisager'),
        ('2 — Bonnes pratiques environnementales', 'mesures_bpe_prises', 'mesures_bpe_envisager'),
        ('3 — Techniques disponibles', 'mesures_tech_prises', 'mesures_tech_envisager'),
        ('4 — Techniques de production plus propres', 'mesures_pp_prises', 'mesures_pp_envisager'),
        ('5 — Gestion préventive et maîtrise des risques', 'mesures_risques_prises', 'mesures_risques_envisager'),
    ]
    for titre, cle_prise, cle_envisager in mesures:
        doc.add_paragraph()
        p = doc.add_paragraph()
        p.add_run(titre).italic = True
        _champ(doc, 'Mesures prises', data.get(cle_prise))
        _champ(doc, 'Mesures à envisager', data.get(cle_envisager))

    for section in doc.sections:
        section.left_margin = Cm(2)
        section.right_margin = Cm(2)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
