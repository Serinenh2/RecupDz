"""
Document Loader — file loading and text chunking only.

Responsibilities:
    - Load files (PDF, DOCX, TXT, Markdown) from disk
    - Split text into overlapping chunks
    - Delegate database sources to DocumentService

Does NOT access Django ORM or import Django models.
Does NOT import repositories directly.
Communicates with DocumentService for all database sources.
"""

from __future__ import annotations

import logging
from dataclasses import dataclass, field
from pathlib import Path
from typing import Any, Dict, List, Optional

from apps.ai_assistant.rag.vector_store import DocumentChunk

logger = logging.getLogger(__name__)


# ---------------------------------------------------------------------------
# Loaded Document
# ---------------------------------------------------------------------------

@dataclass
class LoadedDocument:
    """A loaded document before chunking."""
    text: str
    source: str
    source_type: str
    metadata: Dict[str, Any] = field(default_factory=dict)


# ---------------------------------------------------------------------------
# Text Chunker
# ---------------------------------------------------------------------------

class TextChunker:
    """
    Splits text into overlapping chunks.

    Args:
        chunk_size: Maximum characters per chunk.
        overlap: Number of overlapping characters between chunks.
    """

    def __init__(self, chunk_size: int = 1000, overlap: int = 200) -> None:
        self._chunk_size = chunk_size
        self._overlap = overlap

    def chunk(self, text: str) -> List[str]:
        """Split text into overlapping chunks."""
        if not text or not text.strip():
            return []

        text = text.strip()
        if len(text) <= self._chunk_size:
            return [text]

        chunks = []
        start = 0
        while start < len(text):
            end = start + self._chunk_size

            # Try to break at sentence boundary
            if end < len(text):
                search_start = max(start + self._chunk_size - 200, start)
                last_period = -1
                for punct in ["\n\n", "\n", ". ", "! ", "? ", "؛", "。"]:
                    pos = text.rfind(punct, search_start, end + 100)
                    if pos > last_period:
                        last_period = pos + len(punct)

                if last_period > start:
                    end = last_period

            chunk = text[start:end].strip()
            if chunk:
                chunks.append(chunk)

            start = end - self._overlap
            if start >= len(text):
                break

        return chunks


# ---------------------------------------------------------------------------
# Document Loader
# ---------------------------------------------------------------------------

class DocumentLoader:
    """
    Loads documents from files and chunks them into DocumentChunks.

    For database sources (regulations, waste codes, glossary, procedures),
    delegates to DocumentService.

    Usage:
        loader = DocumentLoader()
        chunks = loader.load_file("regulation.pdf")
        chunks = loader.load_directory("/path/to/docs/")
        chunks = loader.load_from_service(service, ["regulations", "glossary"])
    """

    def __init__(
        self,
        chunk_size: int = 1000,
        chunk_overlap: int = 200,
        service=None,
    ) -> None:
        self._chunker = TextChunker(chunk_size=chunk_size, overlap=chunk_overlap)
        self._service = service

    @property
    def service(self):
        if self._service is None:
            from apps.ai_assistant.services.document_service import DocumentService
            self._service = DocumentService()
        return self._service

    # ------------------------------------------------------------------
    # File loading (pure file I/O)
    # ------------------------------------------------------------------

    def load_file(self, path: str, source_type: Optional[str] = None) -> List[DocumentChunk]:
        """Load a single file and return chunks."""
        file_path = Path(path)
        if not file_path.exists():
            logger.warning("File not found: %s", path)
            return []

        suffix = file_path.suffix.lower()
        if source_type is None:
            source_type = self._detect_type(suffix)

        try:
            if suffix == ".pdf":
                doc = self._load_pdf(str(file_path))
            elif suffix == ".docx":
                doc = self._load_docx(str(file_path))
            elif suffix in (".txt", ".md", ".markdown", ".rst"):
                doc = self._load_text(str(file_path))
            else:
                logger.warning("Unsupported file type: %s", suffix)
                return []

            return self._to_chunks(doc)

        except Exception as exc:
            logger.error("Failed to load %s: %s", path, exc)
            return []

    def load_directory(self, directory: str, extensions: Optional[List[str]] = None) -> List[DocumentChunk]:
        """Load all supported files from a directory."""
        if extensions is None:
            extensions = [".pdf", ".docx", ".txt", ".md"]

        dir_path = Path(directory)
        if not dir_path.exists():
            logger.warning("Directory not found: %s", directory)
            return []

        all_chunks = []
        for ext in extensions:
            for file_path in dir_path.rglob(f"*{ext}"):
                chunks = self.load_file(str(file_path))
                all_chunks.extend(chunks)
                logger.info("Loaded %s: %d chunks", file_path.name, len(chunks))

        return all_chunks

    # ------------------------------------------------------------------
    # Database loading (delegates to DocumentService)
    # ------------------------------------------------------------------

    def load_from_service(
        self,
        sources: Optional[List[str]] = None,
        limit: int = 500,
    ) -> List[DocumentChunk]:
        """
        Load documents from database via DocumentService.

        Args:
            sources: Which sources to load. None = all sources.
            limit: Maximum documents per source.
        """
        try:
            documents = self.service.load_all(sources=sources)

            all_chunks = []
            for doc in documents:
                loaded = LoadedDocument(
                    text=doc["text"],
                    source=doc["source"],
                    source_type=doc["source_type"],
                    metadata=doc.get("metadata", {}),
                )
                all_chunks.extend(self._to_chunks(loaded))

            logger.info("Loaded %d chunks from service", len(all_chunks))
            return all_chunks

        except Exception as exc:
            logger.error("Failed to load from service: %s", exc)
            return []

    def load_regulations(self, limit: int = 500) -> List[DocumentChunk]:
        """Load regulations from database via DocumentService."""
        return self.load_from_service(sources=["regulations"], limit=limit)

    def load_waste_codes(self, limit: int = 1000) -> List[DocumentChunk]:
        """Load waste codes from database via DocumentService."""
        return self.load_from_service(sources=["nomenclature"], limit=limit)

    def load_glossary(self, limit: int = 200) -> List[DocumentChunk]:
        """Load glossary from database via DocumentService."""
        return self.load_from_service(sources=["glossary"], limit=limit)

    def load_procedures(self, limit: int = 200) -> List[DocumentChunk]:
        """Load procedures from database via DocumentService."""
        return self.load_from_service(sources=["procedures"], limit=limit)

    def load_all(self, sources: Optional[List[str]] = None) -> List[DocumentChunk]:
        """Load all knowledge sources via DocumentService."""
        return self.load_from_service(sources=sources)

    # ------------------------------------------------------------------
    # File format loaders (pure file I/O)
    # ------------------------------------------------------------------

    def _load_pdf(self, path: str) -> LoadedDocument:
        """Load a PDF file."""
        try:
            from PyPDF2 import PdfReader
            reader = PdfReader(path)
            text_parts = []
            for page in reader.pages:
                page_text = page.extract_text()
                if page_text:
                    text_parts.append(page_text)

            return LoadedDocument(
                text="\n\n".join(text_parts),
                source=path,
                source_type="pdf",
                metadata={"pages": len(reader.pages)},
            )
        except ImportError:
            logger.error("PyPDF2 not installed")
            return LoadedDocument(text="", source=path, source_type="pdf")

    def _load_docx(self, path: str) -> LoadedDocument:
        """Load a DOCX file."""
        try:
            from docx import Document
            doc = Document(path)
            text_parts = []
            for para in doc.paragraphs:
                if para.text.strip():
                    text_parts.append(para.text)

            return LoadedDocument(
                text="\n\n".join(text_parts),
                source=path,
                source_type="docx",
                metadata={"paragraphs": len(doc.paragraphs)},
            )
        except ImportError:
            logger.error("python-docx not installed")
            return LoadedDocument(text="", source=path, source_type="docx")

    def _load_text(self, path: str) -> LoadedDocument:
        """Load a plain text file."""
        with open(path, encoding="utf-8", errors="replace") as f:
            text = f.read()

        return LoadedDocument(
            text=text,
            source=path,
            source_type="txt",
            metadata={"size_bytes": len(text.encode())},
        )

    # ------------------------------------------------------------------
    # Helpers
    # ------------------------------------------------------------------

    def _to_chunks(self, doc: LoadedDocument) -> List[DocumentChunk]:
        """Convert a LoadedDocument into DocumentChunks."""
        if not doc.text or not doc.text.strip():
            return []

        text_chunks = self._chunker.chunk(doc.text)
        chunks = []
        for i, text in enumerate(text_chunks):
            chunks.append(DocumentChunk(
                text=text,
                source=doc.source,
                source_type=doc.source_type,
                chunk_index=i,
                metadata=doc.metadata,
            ))

        return chunks

    def _detect_type(self, suffix: str) -> str:
        """Detect source type from file extension."""
        mapping = {
            ".pdf": "pdf",
            ".docx": "docx",
            ".txt": "txt",
            ".md": "txt",
            ".markdown": "txt",
            ".rst": "txt",
        }
        return mapping.get(suffix, "unknown")
