"""
Génération du document Word (.docx) du Procès-verbal d'incinération —
contenu équivalent au PDF (apps/inspections/generate_pv.py), présenté de
façon simple et imprimable sous Word. Tient sur une seule page A4, avec un
espace réservé pour la signature en bas de page.
"""
import io
from docx import Document
from docx.shared import Pt, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH

MOIS_FR = [
    '', 'janvier', 'février', 'mars', 'avril', 'mai', 'juin',
    'juillet', 'août', 'septembre', 'octobre', 'novembre', 'décembre',
]
DOTS = '.' * 28


def _texte_date(date_str):
    if not date_str:
        return DOTS, DOTS
    try:
        annee, mois, jour = str(date_str).split('-')[:3]
        return f"{int(jour)}", f"{MOIS_FR[int(mois)]} {annee}"
    except Exception:
        return DOTS, DOTS


def _fmt_date(iso):
    if not iso:
        return ''
    parts = str(iso).split('-')
    if len(parts) != 3:
        return str(iso)
    y, m, d = parts
    return f"{d}/{m}/{y}"


def _espace(doc, taille=4):
    """Espacement fin entre sections — moins encombrant qu'un paragraphe vide normal."""
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(taille)
    p.paragraph_format.space_before = Pt(0)
    r = p.add_run('')
    r.font.size = Pt(2)


def _champ(doc, label, valeur, suite=''):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run(f"{label} ")
    v = str(valeur).strip() if valeur else ''
    r = p.add_run(v if v else DOTS)
    if v:
        r.bold = True
    if suite:
        p.add_run(suite)


def generate_pv_docx(data: dict) -> bytes:
    doc = Document()

    style = doc.styles['Normal']
    style.font.size = Pt(11)
    style.paragraph_format.space_after = Pt(4)
    style.paragraph_format.line_spacing = 1.0

    for section in doc.sections:
        section.top_margin = Cm(1.6)
        section.bottom_margin = Cm(1.4)
        section.left_margin = Cm(2.3)
        section.right_margin = Cm(2.3)

    _champ(doc, 'Raison Sociale', data.get('raison_sociale'))
    _champ(doc, "Agrément d'exploitation N° :", data.get('agrement_exploitation'))
    _champ(doc, 'Adresse :', data.get('adresse'))
    _champ(doc, 'RC :', data.get('rc'))
    _champ(doc, 'NIF :', data.get('nif'))
    _champ(doc, 'NIS :', data.get('nis'))
    _champ(doc, 'ART :', data.get('art'))
    _champ(doc, 'Téléphone :', data.get('telephone'))

    _espace(doc, 8)
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(8)
    p.add_run('Objet : Incinération des Déchets Spéciaux (DS) / Déchets Spéciaux Dangereux (DSD)').bold = True

    pv_numero = data.get('pv_numero') or DOTS
    titre = doc.add_paragraph()
    titre.paragraph_format.space_after = Pt(8)
    titre.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = titre.add_run(f"Procès-verbal d'incinération n° {pv_numero} du {_fmt_date(data.get('date_inspection')) or DOTS}")
    r.bold = True
    r.font.size = Pt(13)

    jour, mois_annee = _texte_date(data.get('date_inspection'))
    societe = data.get('raison_sociale') or DOTS
    site = data.get('site_incineration') or data.get('adresse') or DOTS
    doc.add_paragraph(
        f"L'an deux mille, le {jour} du mois de {mois_annee}, il a été procédé par "
        f"la société : {societe} sur son site d'incinération situé à : {site}"
    )

    doc.add_paragraph("à la destruction par procédé d'incinération des déchets suivants :")

    dechets = data.get('dechets') or [{
        'designation': data.get('designation_dechet') or data.get('designation') or '',
        'quantite': data.get('quantite') or '',
    }]
    table = doc.add_table(rows=1, cols=2)
    table.style = 'Table Grid'
    hdr = table.rows[0].cells
    hdr[0].text = 'Désignation'
    hdr[1].text = 'Quantité'
    for cell in hdr:
        cell.paragraphs[0].runs[0].bold = True
    for d in dechets:
        row = table.add_row().cells
        row[0].text = str(d.get('designation') or '')
        row[1].text = str(d.get('quantite') or '')

    _espace(doc, 6)
    qte_totale = data.get('quantite_totale')
    if not qte_totale and data.get('quantite'):
        qte_totale = f"{data.get('quantite')} {data.get('unite_display') or data.get('unite') or ''}".strip()
    _champ(doc, 'Une quantité totale de', qte_totale,
           suite=f" récupérée par {data.get('recuperateur_nom') or DOTS}")
    _champ(doc, 'Sise :', data.get('recuperateur_adresse'))
    _champ(doc, 'Agrément n°', data.get('recuperateur_agrement'),
           suite=f" du {_fmt_date(data.get('recuperateur_agrement_date')) or DOTS}")

    _espace(doc, 6)
    _champ(doc, 'Les déchets proviennent de :', data.get('generateur_nom'))
    _champ(doc, 'sise', data.get('generateur_adresse'))

    _espace(doc, 8)
    doc.add_paragraph(
        'Nous certifions que les déchets susmentionnés ont été intégralement détruits '
        'par incinération conformément à la réglementation en vigueur et aux '
        'prescriptions techniques de notre installation.'
    )
    doc.add_paragraph('Le présent procès-verbal est établi pour servir et valoir ce que de droit.')

    # ── Espace réservé à la signature ─────────────────────────────────────
    _espace(doc, 14)
    sig = doc.add_table(rows=2, cols=2)
    sig.style = 'Table Grid'
    sig.cell(0, 0).text = "Responsable de l'installation"
    sig.cell(0, 1).text = 'Représentant du récupérateur'
    for cell in (sig.cell(0, 0), sig.cell(0, 1)):
        cell.paragraphs[0].runs[0].bold = True
    sig.cell(1, 0).text = 'Date, signature et cachet :'
    sig.cell(1, 1).text = 'Date, signature et cachet :'
    for cell in (sig.cell(1, 0), sig.cell(1, 1)):
        cell.paragraphs[0].paragraph_format.space_after = Pt(0)
        # Lignes vides pour laisser un espace blanc visible où signer/apposer le cachet
        for _ in range(4):
            blank = cell.add_paragraph()
            blank.paragraph_format.space_after = Pt(0)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
