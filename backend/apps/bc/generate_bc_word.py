"""
Génération du document Word (.docx) des documents commerciaux (BC / Proforma / Facture) —
réplique la mise en page du PDF (voir generate_bc.py) : gabarit générique ou gabarit
SARL INDUREX (en-tête vert, bloc référence, footer RC/NIF + badges ISO) selon le
récupérateur émetteur, pour que le .docx et le .pdf rendent le même document.
"""
import io
from docx import Document
from docx.shared import Pt, Cm, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

from .generate_bc import (
    _recuperateur_info, _calc_totaux, _calc_ligne, _fmt_date, _fmt_montant, _fmt_qte,
    montant_en_lettres, _is_indurex, _INDUREX_NOM, _INDUREX_SLOGAN,
)

_GENERIC_GREEN     = RGBColor(0x3B, 0x6D, 0x11)
_INDUREX_GREEN     = RGBColor(0x3C, 0x7A, 0x42)
_INDUREX_GREEN_HEX = '3C7A42'
_GENERIC_GREEN_HEX = '3B6D11'
_LIGHT_GREEN_HEX   = 'EAF3DE'
_WHITE             = RGBColor(0xFF, 0xFF, 0xFF)
_BLACK             = RGBColor(0x00, 0x00, 0x00)

COL = 17  # largeur de contenu en cm — identique à generate_bc.py (COL = 17 * cm)


# ── Helpers OOXML (shading / bordures ne sont pas exposés par l'API haut niveau
#    de python-docx — nécessaire pour reproduire les en-têtes de tableau vertes
#    et le cadre du bloc client du PDF) ──────────────────────────────────────────

def _shade_cell(cell, hex_color):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color)
    tcPr.append(shd)


def _cell_borders(cell, sz=8, color='000000', edges=('top', 'bottom', 'left', 'right')):
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in edges:
        el = OxmlElement(f'w:{edge}')
        el.set(qn('w:val'), 'single')
        el.set(qn('w:sz'), str(sz))
        el.set(qn('w:space'), '0')
        el.set(qn('w:color'), color)
        tcBorders.append(el)
    tcPr.append(tcBorders)


def _set_col_widths(table, widths_cm):
    table.autofit = False
    for row in table.rows:
        for idx, w in enumerate(widths_cm):
            if idx < len(row.cells):
                row.cells[idx].width = Cm(w)
    for idx, w in enumerate(widths_cm):
        if idx < len(table.columns):
            table.columns[idx].width = Cm(w)


def _set_run(run, *, font='Calibri', size=10, bold=False, italic=False, underline=False, color=None):
    run.font.name = font
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic
    run.underline = underline
    if color is not None:
        run.font.color.rgb = color
    return run


def _zero_spacing(p):
    """Neutralise l'espacement avant/après paragraphe du style Word par défaut
    (~8-10pt) — sans quoi le document parait beaucoup plus aéré que le PDF
    (interligne serré, cf. `leading` des ParagraphStyle de generate_bc.py)."""
    pf = p.paragraph_format
    pf.space_before = Pt(0)
    pf.space_after  = Pt(0)
    return p


def _doc_p(doc, text='', *, align=None, **style):
    """Paragraphe de premier niveau (hors cellule) à espacement nul — équivalent
    du Paragraph reportlab ajouté au `story` dans generate_bc.py."""
    p = doc.add_paragraph()
    _zero_spacing(p)
    if align is not None:
        p.alignment = align
    if text:
        _set_run(p.add_run(text), **style)
    return p


def _cell_lines(cell, lines, align=None):
    """Écrit une ou plusieurs lignes stylées dans une cellule : la 1ère réutilise le
    paragraphe vide déjà présent, les suivantes en ajoutent un nouveau chacune.
    lines: [{text, font, size, bold, italic, color}]"""
    for i, spec in enumerate(lines):
        p = cell.paragraphs[0] if i == 0 else cell.add_paragraph()
        _zero_spacing(p)
        p.alignment = spec.get('align', align)
        run = p.add_run(spec.get('text', ''))
        _set_run(run, font=spec.get('font', 'Calibri'), size=spec.get('size', 10),
                  bold=spec.get('bold', False), italic=spec.get('italic', False),
                  underline=spec.get('underline', False), color=spec.get('color'))
    return cell


def _kv_rows(table, rows, *, label_font='Calibri', label_size=9.5, value_size=9.5):
    """Remplit un tableau 2 colonnes [libellé (gras), valeur] — utilisé pour le bloc
    référence, le bloc client et le tableau RC/NIF/NA/NIS."""
    for i, (label, value) in enumerate(rows):
        row = table.rows[i]
        _cell_lines(row.cells[0], [{'text': label, 'font': label_font, 'size': label_size, 'bold': True}])
        _cell_lines(row.cells[1], [{'text': str(value or ''), 'font': label_font, 'size': value_size}])
    return table


def _add_picture_safe(cell_or_par, path, width_cm, height_cm=None):
    try:
        run = cell_or_par.add_run() if hasattr(cell_or_par, 'add_run') else cell_or_par.paragraphs[0].add_run()
        kw = {'width': Cm(width_cm)}
        if height_cm is not None:
            kw['height'] = Cm(height_cm)
        run.add_picture(path, **kw)
        return True
    except Exception:
        return False


def generate_bc_docx(data: dict) -> bytes:
    rec = _recuperateur_info(data)
    if _is_indurex(rec):
        return _generate_bc_docx_indurex(data, rec)
    return _generate_bc_docx_generique(data, rec)


# ── Gabarit générique ───────────────────────────────────────────────────────────

def _generate_bc_docx_generique(data: dict, rec: dict) -> bytes:
    lignes  = data.get('lignes') or []
    tva_pct = float(data.get('tva_pct') or 19)
    type_doc = data.get('type_document')

    doc = Document()
    for section in doc.sections:
        section.left_margin  = Cm(1.5)
        section.right_margin = Cm(1.5)
        section.top_margin   = Cm(1.2)
        section.bottom_margin = Cm(1.2)

    # ── En-tête : logo + raison sociale ────────────────────────────────────────
    entete = doc.add_table(rows=1, cols=2)
    _set_col_widths(entete, [2.5, COL - 2.5])
    if rec['logo_path']:
        _add_picture_safe(entete.rows[0].cells[0], rec['logo_path'], 2.2, 2.2)
    _cell_lines(entete.rows[0].cells[1], [{
        'text': (rec['nom'] or '').upper(), 'font': 'Calibri', 'size': 20,
        'bold': True, 'italic': True, 'color': _GENERIC_GREEN,
    }])

    if rec['agrement_num']:
        _doc_p(doc, f"Agrément N° {rec['agrement_num']} du {rec['agrement_date']}", size=9)
    adresse_ligne = ' '.join(filter(None, [rec['adresse'], rec['code_postal']]))
    if adresse_ligne:
        _doc_p(doc, adresse_ligne, size=9)

    id_table = doc.add_table(rows=2, cols=2)
    _set_col_widths(id_table, [COL / 2, COL / 2])
    _cell_lines(id_table.rows[0].cells[0], [{'text': f"RC {rec['rc']}", 'size': 9}])
    _cell_lines(id_table.rows[0].cells[1], [{'text': f"NIF {rec['nif']}", 'size': 9}])
    _cell_lines(id_table.rows[1].cells[0], [{'text': f"NA {rec['na']}", 'size': 9}])
    _cell_lines(id_table.rows[1].cells[1], [{'text': f"NIS {rec['nis']}", 'size': 9}])
    _doc_p(doc)

    # ── Date / lieu ─────────────────────────────────────────────────────────────
    _doc_p(doc, f"{rec['commune']} le : {_fmt_date(data.get('date_commande', ''))}",
           align=WD_ALIGN_PARAGRAPH.RIGHT, size=9.5)
    _doc_p(doc)

    # ── Client ──────────────────────────────────────────────────────────────────
    p1 = _doc_p(doc, 'Nome de Client : ', size=9.5)
    _set_run(p1.add_run(str(data.get('client_nom') or '')), size=9.5, bold=True)
    p2 = _doc_p(doc, 'Adresse : ', size=9.5)
    _set_run(p2.add_run(str(data.get('client_adresse') or '')), size=9.5, bold=True)
    _doc_p(doc)

    # ── Titre ───────────────────────────────────────────────────────────────────
    titre_txt = {'PROFORMA': 'Proforma', 'FACTURE': 'Facture'}.get(type_doc, 'Bon de commande')
    titre_tbl = doc.add_table(rows=1, cols=1)
    _set_col_widths(titre_tbl, [8])
    _cell_borders(titre_tbl.rows[0].cells[0])
    _cell_lines(titre_tbl.rows[0].cells[0], [{
        'text': titre_txt, 'size': 13, 'bold': True, 'italic': True, 'align': WD_ALIGN_PARAGRAPH.CENTER,
    }])
    _doc_p(doc)

    # ── Tableau des déchets ──────────────────────────────────────────────────────
    col_w   = [1.2, 6, 2.3, 2, 2.5, 3]
    headers = ['N°', 'Description (Nature des déchets)', 'Quantités', 'Unités', 'Prix unitaires', 'Total HT']
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, col_w)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, _GENERIC_GREEN_HEX)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cell.paragraphs[0].add_run(h), size=9, bold=True, color=_WHITE)

    for i, l in enumerate(lignes, start=1):
        try:
            qte = float(l.get('quantite') or 0)
            pu  = float(l.get('prix_unitaire') or 0)
            ht  = qte * pu
        except (TypeError, ValueError):
            ht = 0.0
        row = tbl.add_row().cells
        vals = [
            str(i), str(l.get('description', '')), str(l.get('quantite', '')), str(l.get('unite', 'KG')),
            f"{pu:,.2f} DZ".replace(',', ' '), f"{ht:,.2f} DZ".replace(',', ' '),
        ]
        for j, val in enumerate(vals):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if j >= 4 else WD_ALIGN_PARAGRAPH.CENTER
            _set_run(row[j].paragraphs[0].add_run(val), size=9)

    _doc_p(doc)

    # ── Récapitulatif HT / TVA / TTC ────────────────────────────────────────────
    total_ht, tva, total_ttc = _calc_totaux(lignes, tva_pct)
    recap = doc.add_table(rows=3, cols=2)
    recap.style = 'Table Grid'
    _set_col_widths(recap, [4, 3])
    recap.alignment = None
    rows_spec = [
        ('Total HT', f"{total_ht:,.2f} DZ", False),
        (f'TVA ({tva_pct:.0f}%)', f"{tva:,.2f} DZ", False),
        ('Total TTC', f"{total_ttc:,.2f} DZ", True),
    ]
    for i, (lbl, val, is_bold) in enumerate(rows_spec):
        row = recap.rows[i]
        _cell_lines(row.cells[0], [{'text': lbl, 'size': 9.5, 'bold': is_bold}])
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run(row.cells[1].paragraphs[0].add_run(val.replace(',', ' ')), size=9, bold=is_bold)
        if is_bold:
            _shade_cell(row.cells[0], _LIGHT_GREEN_HEX)
            _shade_cell(row.cells[1], _LIGHT_GREEN_HEX)
    # Aligner le tableau récap à droite de la page
    tblPr = recap._tbl.tblPr
    jc = OxmlElement('w:jc'); jc.set(qn('w:val'), 'right')
    tblPr.append(jc)

    _doc_p(doc)

    # ── Arrêté ──────────────────────────────────────────────────────────────────
    arrete_nom = {'PROFORMA': 'présente proforma', 'BC': 'présente commande'}.get(type_doc, 'présente facture')
    p_arrete = _doc_p(doc, f"Arrêter la {arrete_nom} en toutes taxes comprises a la somme de : ",
                       size=9, italic=True, underline=True)
    _set_run(p_arrete.add_run(f"{total_ttc:,.2f} DZ".replace(',', ' ')), size=9, italic=True, bold=True)
    validite_jours = data.get('validite_offre_jours')
    if type_doc == 'PROFORMA' and validite_jours not in (None, ''):
        p_arrete.add_run().add_break(WD_BREAK.LINE)
        _set_run(p_arrete.add_run("Validité de l'offre : "), size=9, italic=True, bold=True)
        _set_run(p_arrete.add_run(f"{int(validite_jours)} jours"), size=9, italic=True)

    _doc_p(doc)
    _doc_p(doc)

    # ── Signature ───────────────────────────────────────────────────────────────
    if rec['cachet_path'] or rec['signature_path']:
        sign_p = _doc_p(doc, align=WD_ALIGN_PARAGRAPH.RIGHT)
        if rec['cachet_path']:
            _add_picture_safe(sign_p, rec['cachet_path'], 2.8)
        if rec['signature_path']:
            sign_p.add_run('   ')
            _add_picture_safe(sign_p, rec['signature_path'], 3, 1.6)
    _doc_p(doc, 'Le Gérant', align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)
    if rec['responsable']:
        _doc_p(doc, rec['responsable'], align=WD_ALIGN_PARAGRAPH.RIGHT, size=10)

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()


# ── Gabarit SARL INDUREX ────────────────────────────────────────────────────────

def _generate_bc_docx_indurex(data: dict, rec: dict) -> bytes:
    def v(key, default=''):
        val = data.get(key, default)
        return str(val) if val not in (None, '') else default

    lignes   = data.get('lignes') or []
    tva_pct  = float(data.get('tva_pct') or 19)
    type_doc = data.get('type_document')
    is_proforma = type_doc == 'PROFORMA'
    is_facture  = type_doc == 'FACTURE'
    total_ht, total_tva, total_ttc = _calc_totaux(lignes, tva_pct)

    doc = Document()
    for section in doc.sections:
        section.left_margin   = Cm(1.5)
        section.right_margin  = Cm(1.5)
        section.top_margin    = Cm(1)
        section.bottom_margin = Cm(1.5)

    # ── En-tête : logo + raison sociale + slogan | bloc référence ──────────────
    entete = doc.add_table(rows=1, cols=3)
    _set_col_widths(entete, [2.3, COL - 2.3 - 6.7, 6.7])
    if rec['logo_path']:
        _add_picture_safe(entete.rows[0].cells[0], rec['logo_path'], 2, 2)

    nom_cell = entete.rows[0].cells[1]
    _cell_lines(nom_cell, [
        {'text': _INDUREX_NOM, 'font': 'Calibri', 'size': 20, 'bold': True, 'color': _INDUREX_GREEN},
        {'text': _INDUREX_SLOGAN, 'font': 'Calibri', 'size': 9.5, 'bold': True, 'color': _INDUREX_GREEN},
    ])

    ref_cell = entete.rows[0].cells[2]
    _cell_borders(ref_cell)
    ref_tbl = ref_cell.add_table(rows=5, cols=2)
    _set_col_widths(ref_tbl, [2.6, 4.1])
    _kv_rows(ref_tbl, [
        ('Référence', v('numero')),
        ('Date',      _fmt_date(v('date_commande'))),
        ('Montant',   _fmt_montant(total_ttc)),
        ('Client',    v('ref_client') or v('client_nom')),
        ('Echéance',  _fmt_date(v('date_echeance'))),
    ], label_size=9.5, value_size=9.5)

    _doc_p(doc)

    # ── Titre ───────────────────────────────────────────────────────────────────
    titre_nom = {'PROFORMA': 'Proforma', 'FACTURE': 'Facture'}.get(type_doc, 'Bon de Commande')
    titre_tbl = doc.add_table(rows=1, cols=1)
    _set_col_widths(titre_tbl, [COL])
    titre_cell = titre_tbl.rows[0].cells[0]
    _shade_cell(titre_cell, _INDUREX_GREEN_HEX)
    _cell_lines(titre_cell, [{
        'text': f"{titre_nom} N°: {v('numero')}", 'size': 16, 'bold': True, 'color': _WHITE,
        'align': WD_ALIGN_PARAGRAPH.CENTER,
    }])

    _doc_p(doc)

    # ── Bloc client ──────────────────────────────────────────────────────────────
    client_lignes = [
        ('Réf Client',   v('ref_client')),
        ('N° RC',        v('client_rc')),
        ('NIF',          v('client_nif')),
        ('N° Article',   v('client_numero_article')),
        ('N° I.S',       v('client_nis')),
        ('Tél',          v('client_telephone')),
        ('Fax',          v('client_fax')),
        ('Email',        v('client_email')),
        ('Pièces Liées', v('pieces_liees')),
    ]
    if is_facture:
        client_lignes += [
            ('Mode Paiement', v('mode_paiement')),
            ('Référence',     v('reference_paiement')),
        ]

    bloc_client = doc.add_table(rows=1, cols=2)
    _set_col_widths(bloc_client, [8.5, 8.5])
    gauche_cell = bloc_client.rows[0].cells[0]
    gauche_tbl  = gauche_cell.add_table(rows=len(client_lignes), cols=2)
    _set_col_widths(gauche_tbl, [3.5, 5])
    _kv_rows(gauche_tbl, client_lignes, label_size=10.5, value_size=10.5)

    # Nichée dans sa propre table (plutôt que bordée directement sur la cellule de
    # bloc_client) pour que le cadre n'épouse que la hauteur de son propre contenu
    # (nom + adresse), au lieu de s'étirer sur toute la hauteur de la ligne
    # gauche (9 libellés) — comme le fait le Table flowable indépendant du PDF.
    droite_outer = bloc_client.rows[0].cells[1]
    droite_tbl = droite_outer.add_table(rows=1, cols=1)
    _set_col_widths(droite_tbl, [8.5])
    droite_cell = droite_tbl.rows[0].cells[0]
    _cell_borders(droite_cell)
    droite_lines = [{'text': v('client_nom'), 'font': 'Calibri', 'size': 10.5, 'bold': True}]
    for ligne_adresse in (v('client_adresse') or '').split('\n'):
        if ligne_adresse.strip():
            droite_lines.append({'text': ligne_adresse.strip(), 'size': 10.5})
    _cell_lines(droite_cell, droite_lines)

    _doc_p(doc)

    # ── Tableau des articles ─────────────────────────────────────────────────────
    col_w   = [2.6, 3.3, 1.6, 1.9, 2.1, 2.6, 1.3, 1.6]
    headers = ['Réf Article', 'Désignation', 'Unité', 'Quantité', 'Prix U HT', 'Montant HT', 'R.%', 'Tva%']
    tbl = doc.add_table(rows=1, cols=len(headers))
    tbl.style = 'Table Grid'
    _set_col_widths(tbl, col_w)
    for i, h in enumerate(headers):
        cell = tbl.rows[0].cells[i]
        _shade_cell(cell, _INDUREX_GREEN_HEX)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=_WHITE)

    for l in lignes:
        c = _calc_ligne(l, tva_pct)
        row = tbl.add_row().cells
        vals = [
            str(l.get('ref_article', '')), str(l.get('description', '')), str(l.get('unite', 'KG')),
            _fmt_qte(l.get('quantite')), _fmt_montant(l.get('prix_unitaire')), _fmt_montant(c['ht']),
            f"{c['remise_pct']:.2f}", f"{c['tva_pct']:.2f}",
        ]
        for j, val in enumerate(vals):
            row[j].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT if j in (4, 5) else WD_ALIGN_PARAGRAPH.CENTER
            _set_run(row[j].paragraphs[0].add_run(val), size=10)

    _doc_p(doc)

    # ── Récapitulatifs (gauche : Montant HT/TVA/Montant TVA — droite : totaux) ─
    recap_wrapper = doc.add_table(rows=1, cols=2)
    _set_col_widths(recap_wrapper, [COL - 7, 7])

    g_cell = recap_wrapper.rows[0].cells[0]
    recap_g = g_cell.add_table(rows=2, cols=3)
    recap_g.style = 'Table Grid'
    _set_col_widths(recap_g, [COL - 7 - 6.1, 2.6, 3.5])
    for i, h in enumerate(['Montant HT', 'TVA', 'Montant TVA']):
        cell = recap_g.rows[0].cells[i]
        _shade_cell(cell, _INDUREX_GREEN_HEX)
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cell.paragraphs[0].add_run(h), size=10, bold=True, color=_WHITE)
    for i, val in enumerate([_fmt_montant(total_ht), f"{tva_pct:.2f}", _fmt_montant(total_tva)]):
        cell = recap_g.rows[1].cells[i]
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _set_run(cell.paragraphs[0].add_run(val), size=10)

    d_cell = recap_wrapper.rows[0].cells[1]
    recap_d = d_cell.add_table(rows=5, cols=2)
    recap_d.style = 'Table Grid'
    _set_col_widths(recap_d, [3.5, 3.5])
    d_rows = [
        ('TOTAL H.T',   _fmt_montant(total_ht)),
        ('TOTAL T.V.A', _fmt_montant(total_tva)),
        ('TOTAL T.T.C', _fmt_montant(total_ttc)),
        ('', ''),
        ('NET A PAYER', _fmt_montant(total_ttc)),
    ]
    for i, (lbl, val) in enumerate(d_rows):
        row = recap_d.rows[i]
        _cell_lines(row.cells[0], [{'text': lbl, 'size': 10.5, 'bold': True}])
        row.cells[1].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.RIGHT
        _set_run(row.cells[1].paragraphs[0].add_run(val), size=10.5, bold=True)

    _doc_p(doc)

    # ── Arrêté (montant en lettres + validité de l'offre) ──────────────────────
    arrete_doc_nom = {'PROFORMA': 'la Présente Proforma', 'FACTURE': 'la Présente Facture'}.get(
        type_doc, 'le Présent Bon de Commande')
    p_arrete = _doc_p(doc, f"Arrêtée {arrete_doc_nom} à la Somme de : ", size=10.5)
    _set_run(p_arrete.add_run(montant_en_lettres(total_ttc)), size=10.5, bold=True)
    validite_jours = data.get('validite_offre_jours')
    if is_proforma and validite_jours not in (None, ''):
        p_arrete.add_run().add_break(WD_BREAK.LINE)
        _set_run(p_arrete.add_run("Validité de l'offre : "), size=10.5, bold=True)
        _set_run(p_arrete.add_run(f"{int(validite_jours)} jours"), size=10.5)

    # ── Signature / cachet électroniques ────────────────────────────────────────
    if rec['cachet_path'] or rec['signature_path']:
        _doc_p(doc)
        sign_p = _doc_p(doc, align=WD_ALIGN_PARAGRAPH.CENTER)
        if rec['cachet_path']:
            _add_picture_safe(sign_p, rec['cachet_path'], 4)
        if rec['signature_path']:
            sign_p.add_run('   ')
            _add_picture_safe(sign_p, rec['signature_path'], 4.5, 2.4)

    # ── Pied de page : identité RC/NIF/NIS + badges ISO ─────────────────────────
    footer = doc.sections[0].footer
    footer_p = footer.paragraphs[0]
    footer_p.text = ''
    _zero_spacing(footer_p)
    _set_run(footer_p.add_run(
        f"RC: {rec['rc']}  NIF: {rec['nif']}  Al: {rec['na']}  NIS: {rec['nis']}"
    ), size=7)
    if rec['adresse']:
        p_adr = _zero_spacing(footer.add_paragraph())
        _set_run(p_adr.add_run(rec['adresse']), size=7)
    footer_extra = '  '.join(filter(None, [
        rec['commune'],
        f"Email: {rec['email']}" if rec['email'] else '',
        f"Tél: {rec['telephone']}" if rec['telephone'] else '',
        f"Fax: {rec['fax']}" if rec['fax'] else '',
    ]))
    if footer_extra:
        p_extra = _zero_spacing(footer.add_paragraph())
        _set_run(p_extra.add_run(footer_extra), size=7)

    iso_paths = [p for p in (rec.get('iso_9001_path'), rec.get('iso_14001_path'), rec.get('iso_45001_path')) if p]
    if iso_paths:
        p_iso = _zero_spacing(footer.add_paragraph())
        p_iso.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        for path in iso_paths:
            _add_picture_safe(p_iso, path, 1.5)
            p_iso.add_run('  ')

    buffer = io.BytesIO()
    doc.save(buffer)
    buffer.seek(0)
    return buffer.read()
